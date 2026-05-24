from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("outputs/resume-supplement-20260525")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_FILL = "F2F4F7"
BORDER = "DADCE0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for side, value in values.items():
        node = mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)


def set_run_font(run, name, east_asia=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if east_asia:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def configure_document(doc, lang):
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    base_font = "Calibri" if lang == "en" else "Microsoft JhengHei"
    east_asia = "Microsoft JhengHei" if lang == "zh" else None

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = base_font
    normal.font.size = Pt(10)
    if east_asia:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in [
        ("Heading 1", 13.5, BLUE, 10, 4),
        ("Heading 2", 11.5, BLUE, 6, 2),
        ("Heading 3", 10.5, DARK_BLUE, 4, 2),
    ]:
        style = styles[style_name]
        style.font.name = base_font
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        if east_asia:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Bullet 2"]:
        style = styles[style_name]
        style.font.name = base_font
        style.font.size = Pt(9.6)
        if east_asia:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.14)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.08


def add_title(doc, name, tagline, contact, lang):
    base_font = "Calibri" if lang == "en" else "Microsoft JhengHei"
    east_asia = "Microsoft JhengHei" if lang == "zh" else None

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(name)
    set_run_font(r, base_font, east_asia)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(tagline)
    set_run_font(r, base_font, east_asia)
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(contact)
    set_run_font(r, base_font, east_asia)
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED

    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D9E2F3")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def add_section(doc, title):
    p = doc.add_heading(title, level=1)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)


def add_skill_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [1900, 7460])
    hdr = table.rows[0].cells
    hdr[0].text = rows[0][0]
    hdr[1].text = rows[0][1]
    for cell in hdr:
        set_cell_shading(cell, LIGHT_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.4)
    for label, detail in rows[1:]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        for p in cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.2)
                r.font.color.rgb = DARK_BLUE
        for p in cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9.2)
    set_table_width(table, [1900, 7460])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_role(doc, title, company, dates, bullets, lang):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p.add_run(" | " if lang == "en" else "，")
    p.add_run(company)
    p.add_run(" | " if lang == "en" else "，")
    r = p.add_run(dates)
    r.italic = True
    r.font.color.rgb = MUTED
    add_bullets(doc, bullets)


def build_english():
    doc = Document()
    configure_document(doc, "en")
    add_title(
        doc,
        "Chang Wei-Lin",
        "Algorithm Engineer | Audio, Image, and Physiological Signal Processing | Embedded C/C++ and Python | AI and Deep Learning",
        "Hsinchu, Taiwan | x111281@gmail.com",
        "en",
    )

    add_section(doc, "Professional Summary")
    add_para(
        doc,
        "Algorithm engineer with a physics M.S. and 10+ years across embedded audio, signal processing, thermal vision, physiological sensing, RF test automation, and optical microscopy R&D. Strong at turning physical signals into robust features and production-ready algorithms: prototype in Python, validate with data, then implement optimized C/C++ or Assembly on resource-constrained MCU/DSP platforms.",
    )
    add_para(
        doc,
        "Target roles: Artificial Intelligence Engineer, Deep Learning Engineer, Algorithm Engineer, Software Engineer."
    )

    add_section(doc, "Core Strengths")
    add_bullets(
        doc,
        [
            ("Signal and algorithm development: ", "audio coding, pitch tracking, time-scale modification, VAD/noise reduction, feedback suppression, beat/syllable segmentation, image scaling, and image compression."),
            ("Embedded implementation: ", "C/C++/Assembly on MCU and Bluetooth chip platforms, with fixed-point and floating-point optimization for limited compute and memory."),
            ("AI/ML and computer vision: ", "OpenCV, YOLO-based posture/fall detection, thermal imaging, microwave radar features, Random Forest, CNN, DNN, TensorFlow, and PyTorch."),
            ("Physics-grounded modeling: ", "optics, microscopy, femtosecond laser systems, nanoparticle tracking, physiological signals, and physics-based data analysis."),
            ("AI-assisted productivity: ", "active use of Gemini, ChatGPT, Claude, Grok, Antigravity, and Claude Code for coding, verification, learning, and technology tracking."),
        ],
    )

    add_section(doc, "Technical Skills")
    add_skill_table(
        doc,
        [
            ("Category", "Details"),
            ("Programming", "C/C++ (Qt, OpenCV), Python (TensorFlow, PyTorch), C#, Assembly, Linux/Bash, MATLAB, LabVIEW"),
            ("Algorithms", "Audio signal processing; image and thermal signal processing; physiological and physical signal processing; ML/DL; web scraping"),
            ("Physics / Optics", "Tunable femtosecond pulsed lasers; confocal microscopy; two-photon excitation microscopy; harmonic imaging microscopy; interferometric/nanoparticle scattering microscopy; Newtonian telescope"),
        ],
    )

    add_section(doc, "Experience")
    add_role(
        doc,
        "Algorithm Engineer",
        "NyQuest Technology Co. Ltd.",
        "Nov 2022 - Present",
        [
            "Develop and optimize speech algorithm libraries for core products including voice-control ICs, voice microcontrollers, and LCD microcontrollers.",
            "Prototype and verify algorithms in Python, then implement high-performance C routines for resource-constrained MCU environments.",
            "Implemented and improved ADPCM/SBC audio coding, pitch tracking, time-stretching and pitch-shifting with resampling/WSOLA, syllable segmentation, feedback suppression, beat tracking, image scaling, and image compression.",
            "Use AI-assisted coding and verification workflows where appropriate to accelerate implementation, testing, and technical learning.",
        ],
        "en",
    )
    add_role(
        doc,
        "Algorithm Engineer",
        "ADE Technology Inc.",
        "Jul 2020 - Oct 2022",
        [
            "Participated in key technology development integrating thermal imaging cameras and microwave radar for healthcare, epidemic prevention, and livestock applications.",
            "Developed non-contact respiration and heartbeat monitoring algorithms using Python simulation and C++ implementation.",
            "Built a microwave marker GUI to improve data labeling, usability, and visualization for algorithm development.",
            "Trained Random Forest, CNN, and DNN models with microwave data for symptom recognition.",
            "Developed thermal-imaging posture and fall detection using OpenCV and YOLO.",
            "Automated remote Linux device control and data collection with Bash; used SMAC stepper motors to simulate respiration and cardiac motion for algorithm debugging.",
        ],
        "en",
    )
    add_role(
        doc,
        "Software Engineer",
        "UnlimiterHear Co., Ltd.",
        "May 2016 - Jun 2020",
        [
            "Worked on algorithm R&D for Bluetooth headsets and hearing aids, focusing on hearing-health applications.",
            "Simulated algorithms in Python and implemented production solutions in C and Assembly on CSR8670 and AB1539/T Bluetooth chips.",
            "Developed and optimized voice enhancement algorithms including VAD, likelihood ratio test, audio noise reduction, and sub-bass frequency shifting.",
            "Implemented hearing-protection features including EQ, dynamic range compression, and binaural beat generation.",
            "Generated natural sound effects using randomization, white noise, and droplet physics; optimized floating-point and fixed-point DSP arithmetic for resource-constrained platforms.",
        ],
        "en",
    )
    add_role(
        doc,
        "RF Engineer, Multimedia BU",
        "Quanta Computer Inc.",
        "Nov 2014 - May 2016",
        [
            "Supported client RF component development through testing, pilot production, and result reporting.",
            "Developed and maintained automated test programs for instrument control and data analysis.",
            "Collected production-line data, analyzed failure patterns, and proposed yield-improvement actions.",
        ],
        "en",
    )
    add_role(
        doc,
        "Research Assistant",
        "Institute of Atomic and Molecular Sciences, Academia Sinica",
        "Jul 2013 - Apr 2014",
        [
            "Joined Prof. Chia-Lung Hsieh's lab after military service and designed an interferometric scattering optical microscope (iSCAT) for nanoparticle imaging on biological membranes.",
            "Used high-speed cameras (1k-100k fps) to localize nanoparticles on cell membranes and track dynamic trajectories.",
            "Analyzed large-scale trajectory data to study physical properties of cell membranes; contributed to a publication in Optics Express.",
        ],
        "en",
    )

    add_section(doc, "Patents and Publications")
    add_bullets(
        doc,
        [
            "Patent applications with colleagues: Sound adjustment device for hearing protection and sound adjustment method thereof; Sound adjustment method and sound adjustment device; Sound playback device and method for masking interference sound through masking noise signal thereof.",
            "Publication: Shot-noise limited localization of single 20 nm gold particles with nanometer spatial precision within microseconds, Optics Express.",
        ],
    )

    add_section(doc, "Education")
    add_role(
        doc,
        "Master, Institute of Physics",
        "National Taiwan University",
        "Sep 2009 - Jun 2012",
        [
            "Joined Prof. Shi-Wei Chu's Biomedical Optics Lab; studied nonlinear crystals, photonic crystal fibers, and femtosecond pulsed laser sources.",
            "Developed super-resolution fluorescence microscopy and nonlinear optical microscopy for non-invasive 3D biomedical imaging.",
            "Collaborated with Academia Sinica, NTHU, NYMU, FCU, and Osaka University; served as a teaching assistant for undergraduate optics and physics courses.",
            "Thesis: Tunable Pulse Laser Generation from 1.3 μm to 1.8 μm by Periodically Poled Lithium Niobate.",
        ],
        "en",
    )
    add_role(
        doc,
        "Bachelor, Department of Physics",
        "National Cheng Kung University",
        "Sep 2005 - Jun 2009",
        [
            "Participated in a plasma detection project at the Plasma and Space Science Center, focusing on space instrument development.",
            "Simulated antenna and transmission-line structures, calculated detector spectra and transmission efficiency, and observed plasma properties through high-frequency microwave interactions.",
            "Active in student groups including the Departmental Association, Badminton Team, Yilan Alumni Club, Hip-Hop Dance Club, and Badminton Club.",
        ],
        "en",
    )

    add_section(doc, "Continuous Learning and Interests")
    add_para(
        doc,
        "Avid reader and hands-on learner who studies algorithms, deep learning, modern C++, design patterns, decision science, and emerging AI systems. Enjoys hiking, strength training, and board games."
    )
    return doc


def build_chinese():
    doc = Document()
    configure_document(doc, "zh")
    add_title(
        doc,
        "張維麟",
        "演算法工程師 | 音訊、影像與生理訊號處理 | Embedded C/C++ 與 Python | AI 與深度學習",
        "Hsinchu, Taiwan | x111281@gmail.com",
        "zh",
    )

    add_section(doc, "專業摘要")
    add_para(
        doc,
        "具物理碩士背景的演算法工程師，擁有十年以上演算法、軟體與研發經驗，橫跨嵌入式音訊、訊號處理、熱影像、生理感測、射頻測試自動化與光學顯微研發。擅長從物理訊號與資料中萃取關鍵特徵，以 Python 建模與驗證，再以 C/C++ 或 Assembly 在資源有限的 MCU/DSP 平台上完成高效率實作。"
    )
    add_para(doc, "投遞職務：人工智慧工程師、深度學習工程師、演算法工程師、軟體工程師。")

    add_section(doc, "核心能力")
    add_bullets(
        doc,
        [
            ("訊號與演算法開發：", "音訊編碼、音高追蹤、變速變調、VAD/降噪、迴授抑制、節拍/音節分割、圖像縮放與圖像壓縮。"),
            ("嵌入式實作：", "在 MCU 與藍牙晶片平台上以 C/C++/Assembly 實作演算法，並進行定點數、浮點數與資源受限環境下的效能優化。"),
            ("AI/ML 與電腦視覺：", "OpenCV、YOLO 姿態/跌倒偵測、熱影像、微波雷達特徵、Random Forest、CNN、DNN、TensorFlow 與 PyTorch。"),
            ("物理建模能力：", "光學、顯微鏡、飛秒脈衝雷射、奈米粒子追蹤、生理訊號與物理訊號分析。"),
            ("AI 輔助生產力：", "持續使用 Gemini、ChatGPT、Claude、Grok、Antigravity、Claude Code 協助編碼、驗證、學習新知與追蹤前沿技術。"),
        ],
    )

    add_section(doc, "技術技能")
    add_skill_table(
        doc,
        [
            ("類別", "內容"),
            ("程式語言", "C/C++ (Qt, OpenCV)、Python (TensorFlow, PyTorch)、C#、Assembly、Linux/Bash、MATLAB、LabVIEW"),
            ("演算法", "音訊訊號處理、影像與熱影像訊號處理、生理與物理訊號處理、ML/DL、網路爬蟲"),
            ("物理 / 光學", "可調頻飛秒脈衝雷射、共軛焦顯微鏡、雙光子激發顯微鏡、倍頻顯微鏡、干涉式/奈米粒子散射顯微鏡、牛頓式天文望遠鏡"),
        ],
    )

    add_section(doc, "工作經歷")
    add_role(
        doc,
        "演算法工程師",
        "九齊科技",
        "Nov 2022 - 至今",
        [
            "參與開發與優化應用於語音控制 IC、語音微控制器與 LCD 微控制器等核心產品上的語音演算法函式庫。",
            "以 Python 進行演算法模擬與驗證，再以 C 語言在資源有限的 MCU 環境中完成高效率實作。",
            "實作與改進 ADPCM/SBC 音訊編碼、音高追蹤、resampling/WSOLA 變速變調、音節分割、迴授抑制、節拍偵測、圖像縮放與圖像壓縮。",
            "在合適情境導入 AI 輔助編碼與驗證流程，加速實作、測試與技術學習。",
        ],
        "zh",
    )
    add_role(
        doc,
        "演算法工程師",
        "亞迪電子",
        "Jul 2020 - Oct 2022",
        [
            "參與結合熱像鏡頭與微波雷達的關鍵技術開發，應用於醫療照護、防疫與畜牧領域。",
            "以 Python 模擬、C++ 實作非接觸式呼吸與心跳監測演算法。",
            "設計微波標記 GUI 工具，提高資料標註、操作便利性與視覺化能力。",
            "使用微波資料訓練 Random Forest、CNN 與 DNN 模型，進行病徵辨識。",
            "運用 OpenCV 與 YOLO 開發熱影像姿態與跌倒辨識系統。",
            "以 Bash 控制遠端 Linux 裝置並自動收集資料；使用 SMAC 步進馬達模擬病人呼吸與心跳運動，支援演算法調試。",
        ],
        "zh",
    )
    add_role(
        doc,
        "軟體工程師",
        "元鼎音訊",
        "May 2016 - Jun 2020",
        [
            "參與藍牙耳機與輔聽器的演算法研發，專注於聽力健康技術。",
            "以 Python 進行演算法模擬，並在 CSR8670、AB1539/T 等藍牙晶片上以 C 與 Assembly 完成產品實作。",
            "開發與優化語音增強演算法，包括 VAD、Likelihood Ratio Test、音訊降噪與超低音移頻技術。",
            "實作 EQ、Dynamic Range Compression 與雙耳拍音生成，提供聽力保護與舒適聽覺體驗。",
            "利用亂數、白噪聲與水滴物理生成自然聲音效果；優化浮點數與定點數 DSP 運算，使系統能在資源受限平台上穩定執行。",
        ],
        "zh",
    )
    add_role(
        doc,
        "射頻工程師",
        "廣達電腦 Multimedia BU",
        "Nov 2014 - May 2016",
        [
            "協助客戶進行射頻元件開發、測試實驗、試產與結果報告。",
            "撰寫與維護自動化測試程式，進行儀器控制與資料分析。",
            "收集產線資料、分析錯誤模式，並提出良率改善方案。",
        ],
        "zh",
    )
    add_role(
        doc,
        "研究助理",
        "中央研究院 原子與分子科學研究所",
        "Jul 2013 - Apr 2014",
        [
            "退伍後加入謝佳龍老師實驗室，設計干涉式散射光學顯微鏡 (iSCAT)，建立生物膜奈米粒子影像。",
            "使用高速攝影機 (1k-100k fps) 對細胞膜上的奈米粒子定位，並追蹤動態軌跡。",
            "分析大量動態軌跡資料以研究細胞膜物理特性，並與團隊將成果發表於 Optics Express。",
        ],
        "zh",
    )

    add_section(doc, "專利與發表")
    add_bullets(
        doc,
        [
            "與同事合作申請專利：具聽力保護功能之聲音調整裝置及其聲音調整之方法；聲音調整方法及聲音調整裝置；聲音播放裝置及其透過遮噪音訊遮蓋干擾音之方法。",
            "期刊論文：Shot-noise limited localization of single 20 nm gold particles with nanometer spatial precision within microseconds, Optics Express.",
        ],
    )

    add_section(doc, "學歷")
    add_role(
        doc,
        "碩士，物理研究所",
        "國立臺灣大學",
        "Sep 2009 - Jun 2012",
        [
            "加入朱士維老師生醫光學實驗室，研究非線性晶體、光子晶體光纖與飛秒脈衝雷射光源。",
            "開發超解析度螢光顯微鏡與非線性光學顯微鏡，應用於非侵入式 3D 生醫影像。",
            "與中研院、清大、陽明、逢甲、大阪大學等單位進行合作計畫；曾擔任光學與物理大學課程助教。",
            "論文題目：Tunable Pulse Laser Generation from 1.3 μm to 1.8 μm by Periodically Poled Lithium Niobate。",
        ],
        "zh",
    )
    add_role(
        doc,
        "學士，物理學系",
        "國立成功大學",
        "Sep 2005 - Jun 2009",
        [
            "加入電漿與太空科學中心的電漿探測計畫，參與太空儀器開發。",
            "透過軟體模擬天線與傳輸線結構，計算探測器頻譜與傳輸效率，並以高頻微波與電漿粒子的交互作用觀察電漿物理特性。",
            "大學期間積極參與系學會、系羽球隊、蘭友會、熱舞社與羽球社。",
        ],
        "zh",
    )

    add_section(doc, "持續學習與興趣")
    add_para(
        doc,
        "喜歡閱讀並對世界保持好奇，持續學習演算法、深度學習、Modern C++、設計模式、判斷與決策，以及最新 AI 技術。興趣包含登山、重訓與桌遊。"
    )
    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = {
        "Chang_Wei-Lin_resume_supplemented_20260525_English.docx": build_english(),
        "張維麟履歷-補充版-20260525-中文.docx": build_chinese(),
    }
    for filename, doc in outputs.items():
        path = OUT_DIR / filename
        doc.save(path)
        print(path.resolve())


if __name__ == "__main__":
    main()
